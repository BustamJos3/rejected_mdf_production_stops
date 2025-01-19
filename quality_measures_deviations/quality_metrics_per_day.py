#!/usr/bin/env python
# coding: utf-8

# # TODO
# * Discriminar datos por tipo de panel (4 tipos*cantidad de calibres en medición)

# In[41]:


import pandas as pd #modules
import numpy as np
import matplotlib
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
import os
import pathlib
from pathlib import Path
from datetime import datetime


# In[42]:


directory = Path('./data') #get current work directory
directory.mkdir(exist_ok=True)
matching_files = list(directory.glob("*obj*.xlsx"))  # Busca archivos que contengan 'obj' y tengan extensión .xlsx
print("Archivos encontrados:", matching_files)


# In[43]:


dict_data_pointer={} #dict to store files as dfs
list_wanted_sheets=["data","control"]
for i in matching_files: # Read the Excel file 
    file_path = str(i)  # Update this with the path
    print(file_path)
    for sheet in list_wanted_sheets:
        if sheet=="control":
            rows_start=5
        else:
            rows_start=6
        df = pd.read_excel(file_path,sheet_name=sheet,skiprows=rows_start)
        df_name=file_path.split("obj_")[1].split(".")[0]+"_"+sheet #split str with "obj_" and the  "." char and take the file name
    #row_numbers_nan = df.index[df[date_col].isna()].tolist() # Get the row numbers where 'Fecha' or 'Fecha Paro' is NaN
    #row_numbers = list(set(row_numbers_nan)) # lists of row with nan numbers
    #filtered_df = df.drop(index=row_numbers) # Filter the DataFrame to keep only the rows that are not in row_numbers
        dict_data_pointer[f"{df_name}"]=df #store filtered df in dictionary data pointer
print(list(dict_data_pointer.keys())) #see keys on dictionary to check callability


# In[44]:


df_measured=dict_data_pointer['quality_data_data']
df_measured


# In[45]:


df_control=dict_data_pointer['quality_data_control']
df_control


# # Plot measured variables and control variables
# * op1: lines

# In[47]:


list_measured_cols=["% resina sólida/fibra seca","Tracción interna (kg/cm2)","Peso manta (kg/m2)",
                    "Velocidad línea (mm/s)","% humedad","% hinchamiento 24 h ",
                   "% Parafina sólida/fibra seca2"] #"Calibre",
df_measured.loc[:,["Fecha",list_measured_cols[-1]]]


# In[48]:


df_measured["Fecha"].unique()


# In[49]:


df_measured["Fecha"]=df_measured["Fecha"].astype("str")
dates=[date for date in df_measured["Fecha"].unique()]
dates


# In[50]:


len(list_measured_cols)


# In[134]:


list_control_cols=["Calibre \n( mm )","Resina solida/fibra seca (%)","Parafina sólida/fibra seca %","Peso manta Promedio (kg/m2)",
                   "Peso manta (kg/m2)","Vel. Linea Mínimo  (mm/s)","TRACCION MINIMA  (Kg/cm^2)"]
df_control.loc[:,[list_control_cols[0],list_control_cols[-1],"Tipo"]]


# In[52]:


col_matches={0:1,1:-1,2:4,3:5,}
dict_measured_control_matches={list_measured_cols[key]:list_control_cols[value] for (key,value) in col_matches.items() }
dict_measured_control_matches


# In[53]:


calibres_measured=[calibre for calibre in df_measured["Calibre"].unique()]
calibres_measured


# In[54]:


type_caliber_measured=[type for type in df_measured["Tipo"].unique()]
type_caliber_measured


# # TODO
# * complete getting of control values for respective data and caliber

# In[56]:


df_control.loc[(df_control[list_control_cols[0]]==calibres_measured[0] ) & (df_control["Tipo"]=="HDF" ) ]


# In[57]:


df_measured_current_calibre=df_measured[df_measured["Calibre"]==calibres_measured[0]]
df_measured_current_calibre


# In[58]:


current_data=df_measured_current_calibre.loc[:,["Fecha",list_measured_cols[0]]]
current_data


# # Plotting
# ## TODO
# * plot the control lines

# In[117]:


current_date=datetime.today()
current_date=current_date.date()
current_date.day


# In[119]:


path=r'./imgs_reports_daily/qlty_deviations_plots/month_{}/day_{}/'.format(11,26)
path


# In[121]:


directory = Path(path) #get current work directory
directory.mkdir(exist_ok=True)


# In[225]:


test=df_control.loc[ (df_control[list_control_cols[0]]==calibres_measured[0])&(df_control[type_mdf_col]==type_caliber_measured[0]) ]
test


# In[236]:


colors = ('Green', 'Red', 'Blue',"Brown")
date_col="Fecha"
marker_styles=['*','s','o','v','X','x'] #markers to 6 plots
line_styles=['-','--','-.',':','-','--']
calibers_col="Calibre"
type_mdf_col="Tipo"
data_cols_pointers=[[0,4],[4,-1]] #set from where to start calling cols and to where stop the calling
titles_of_imgs=[] #list to store title of data just plotted
for type_mdf in type_caliber_measured:
    print(type_mdf)
    for caliber in calibres_measured:
        print(caliber)
        df_measured_current_calibre=df_measured.loc[ (df_measured[calibers_col]==caliber)&(df_measured[type_mdf_col]==type_mdf) ] #filter measurements for current caliber
        if len(df_measured_current_calibre)<1: #check if querie is not empty
            continue
        intAxNo_plot_number = 0 #counter to select aditional y-axis
        fig, axs = plt.subplots(2,sharex=True) #make 2 plot artists
        axis_counter=0 #counter to see when second axs is selected
        for ax,pointer in zip(axs,data_cols_pointers): #run through axis artists for make figure with 2 plots
            #caliber designation: for future, convert to for loop
            data_cols_start_idx=pointer[0]
            data_cols_end_idx=pointer[-1]
            if axis_counter>0:
                axes = [ax, ax.twinx()] #only 2 vars to plot remaining
            else:
                # Twin the x-axis twice to make independent y-axes.
                axes = [ax, ax.twinx(), ax.twinx(),ax.twinx()]
            axis_counter+=1
            # Make some space on the right side for the extra y-axis.
            fig.subplots_adjust(right=0.85)
            # Move the last y-axis spine over to the right by 20% of the width of the axes
            y_axis_displacement=0.25
            for axe in axes[1:]: #for all multiple axes, except for 1st one
                axe.spines['right'].set_position(('axes', -1*y_axis_displacement))
                y_axis_displacement+=0.2 #increase displacement
            # To make the border of the right-most axis visible, we need to turn the frame
            # on. This hides the other plots, however, so we need to turn its fill off.
            axes[-1].set_frame_on(True)
            axes[-1].patch.set_visible(False)
            # And finally we get to plot things...
            current_dates=[date for date in df_measured_current_calibre[date_col].unique()] #dates for current caliber
            intAxNo_selector=0
            labelpad_value=45 #set desired distance for y_label name
            for ax, color,(idx_data,data_col) in zip(axes, colors,enumerate(list_measured_cols[data_cols_start_idx:data_cols_end_idx])):
                intAxNo_selector+=1
                print(data_col)
                current_data=df_measured_current_calibre.loc[:,[date_col,data_col]]
                control_values=df_control.loc[ (df_control[list_control_cols[0]]==caliber)&(df_control[type_mdf_col]==type_mdf) ] #filter by caliber and type on control values
                try: #try to get control values for vars to be plotted
                    current_control_val=control_values.at[ list(control_values.index)[0],dict_measured_control_matches[data_col] ]
                except: #current exception, empty querie
                    current_control_val="Sin dato"
                data = np.array([val[0] for val in current_data.groupby([date_col]).mean().values])
                ax.plot(current_dates,data,label="control={}".format(current_control_val),marker=marker_styles[intAxNo_plot_number], linestyle=line_styles[intAxNo_plot_number],color=color)
                amt_xticks=range(len(current_dates))
                ax.set_xticks(amt_xticks)
                ax.set_xticklabels(current_dates,rotation=90) #rotate x axis labels 90º to be displayed vertically
                if (intAxNo_selector > 1):
                    ax.set_ylabel(f' {data_col}', color=color, labelpad = -1*labelpad_value )
                    ax.get_yaxis().set_tick_params(direction='out')
                    labelpad_value+=5
                else:
                    ax.grid() #activate grid with the 1st y axis
                    ax.set_ylabel(f' {data_col}', color=color, labelpad = +0 ) #1st y-axis
                ax.tick_params(axis='y', colors=color) #adjust y-axis names
                intAxNo_plot_number += 1
            axes[0].set_xlabel('Fecha de medición')
        img_title="Paneles Calibre {} $[mm]$ tipo {}".format(caliber,type_mdf)
        fig.suptitle(img_title)
        fig.legend(loc="lower left")
        img_save_path=img_title.replace(" ","_")+".png"
        print(path+img_save_path)
        titles_of_imgs.append(path+img_save_path) #store title of image to display on report
        plt.savefig(path+img_save_path,bbox_inches='tight')
        plt.show()


# # make word report

# In[238]:


import docx as dx
from docx.shared import Inches, Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


# In[239]:


#--->functions
def aligment_paragraph(p,agliment): #text aligment setter
    '''
    the function sets the aligment for a paragraph object 
    through the string aligment, which can only be the 
    next values:
    -center
    -right
    -left
    -justify
    '''
    str_valid_aligments="center, right, left, justify"
    try:
        a=agliment.upper()
        p_format = p.paragraph_format
        p_format.alignment = eval('WD_ALIGN_PARAGRAPH.'+a)
    except:
        if a.lower() not in str_valid_aligments:
            print(f'is {a} corresponding with')
        else:
            str_other_exceptions="check paragraph object valid state"
            print(str_other_exceptions)


# In[249]:


#---doc creation
doc=dx.Document()
date_today=datetime.today()
current_day=date_today.date()

str_title='Mediciones de calidad, {}'.format(current_day.strftime('%d-%m-%Y'))
level_title=0
title=doc.add_heading(str_title,level_title)    #report title

level_heading1=1
str_objectives_title='Objetivos'
doc.add_heading(str_objectives_title,level_heading1)   #objectives
paragraph = doc.add_paragraph()
paragraph.add_run("\n-Realizar un reporte gráfico de las variables más relevantes medidas en el laboratorio de calidad."+
                 "\n-Mostrar el conjunto de valores medidos para los calibres presentes en el intervalo de tiempo inicio de mes hasta el día del reporte."+
                 "\n-Presentar los valores medidos junto con sus valores de control dados para facilitar la lectura y comprensión de las curvas.")
str_aligment='justify'  #choosen aligment
aligment_paragraph(paragraph,str_aligment)  #setting aligment

str_objectives_title='Curvas'
doc.add_heading(str_objectives_title,level_heading1)   #plots
p = doc.add_paragraph()
for title_of_img in titles_of_imgs:
    r = p.add_run()
    r.add_picture(title_of_img,width=Inches(6.5)) #, height=Inches(.7)

str_objectives_title='Valores de correlación'
doc.add_heading(str_objectives_title,level_heading1)   #plots
p = doc.add_paragraph()
r = p.add_run("\nSe presenta a continuación el reporte de correlaciones lineales para las variables involucradas.")
str_aligment='justify'  #choosen aligment
aligment_paragraph(paragraph,str_aligment)  #setting aligment

str_doc_title='report_qlty_deviations_template.docx'
doc.save(str_doc_title)    #save document


# # The End
